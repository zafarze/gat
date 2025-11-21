# D:\GAT\core\views\student_exams.py (ПОЛНАЯ ИСПРАВЛЕННАЯ ВЕРСИЯ)

import random
import json
from collections import defaultdict
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import JsonResponse, HttpResponse
from django.views.decorators.http import require_POST
from django.template.loader import render_to_string
from django.http import HttpResponse
import weasyprint
from django.conf import settings

# --- Imports for python-docx ---
from docx import Document
from docx.shared import Mm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Импортируем модели ---
from ..models import StudentResult, BankQuestion, Subject, GatTest


# =============================================================================
# --- ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ДЛЯ WORD ---
# =============================================================================

def set_columns(section, cols):
    """
    Устанавливает реальное количество колонок для секции Word через XML.
    Позволяет делать "газетную" верстку с разделительной линией.
    """
    sectPr = section._sectPr
    cols_xml = sectPr.xpath('./w:cols')[0]
    cols_xml.set(qn('w:num'), str(cols))
    cols_xml.set(qn('w:space'), '425') # Расстояние между колонками (около 0.75см)
    
    # ✨ ЭТА СТРОКА ДОБАВЛЯЕТ ЛИНИЮ МЕЖДУ КОЛОНКАМИ
    cols_xml.set(qn('w:sep'), '1')


# =============================================================================
# --- VIEWS ДЛЯ СТУДЕНТОВ ---
# =============================================================================

@login_required
def exam_list_view(request):
    if not hasattr(request.user, 'profile') or request.user.profile.role != 'STUDENT':
        return redirect('core:dashboard')
    
    student = request.user.profile.student
    student_results = StudentResult.objects.filter(student=student).select_related('gat_test').order_by('-gat_test__test_date')

    context = {
        'title': 'Мои экзамены',
        'results': student_results,
    }
    return render(request, 'student_dashboard/exam_list.html', context)


@login_required
def exam_review_view(request, result_id):
    result = get_object_or_404(StudentResult, id=result_id)
    gat_test = result.gat_test
    if not request.user.profile.student == result.student:
        messages.error(request, "У вас нет доступа к этому результату.")
        return redirect('core:student_dashboard')

    questions_by_subject = {}
    all_bank_questions = gat_test.questions.select_related('topic', 'subject', 'school_class').prefetch_related('options').order_by('subject__name', 'id')
    student_answers_dict = result.scores_by_subject or {}

    for bq in all_bank_questions:
        subject_name = bq.subject.name
        if subject_name not in questions_by_subject: questions_by_subject[subject_name] = []
        subject_answers = student_answers_dict.get(str(bq.subject_id), {})
        was_correct = subject_answers.get(str(bq.id))
        questions_by_subject[subject_name].append({
            'question': bq, 'student_was_correct': was_correct, 'options': list(bq.options.all())
        })

    return render(request, 'student_dashboard/exam_review.html', {'title': f'Разбор теста: {gat_test.name}', 'result': result, 'questions_by_subject': questions_by_subject})


# =============================================================================
# --- PREVIEW БУКЛЕТА (HTML) ---
# =============================================================================

@login_required
def gat_test_booklet_preview(request, test_pk):
    test = get_object_or_404(
        GatTest.objects.prefetch_related('questions__options', 'questions__subject'), 
        pk=test_pk
    )
    
    raw_questions = list(test.questions.all())
    
    # 1. Создаем карту порядка для вопросов
    order_map = {qid: idx for idx, qid in enumerate(test.question_order or [])}

    # 2. Группируем вопросы по предметам
    questions_by_subject = defaultdict(list)
    for q in raw_questions:
        questions_by_subject[q.subject].append(q)

    # 3. Определяем порядок ПРЕДМЕТОВ
    subject_order_map = {}
    for subject, q_list in questions_by_subject.items():
        if q_list:
            min_index = min([order_map.get(q.id, 999999) for q in q_list])
        else:
            min_index = 999999
        subject_order_map[subject] = min_index

    # 4. Сортируем предметы
    sorted_subjects = sorted(
        questions_by_subject.keys(),
        key=lambda s: (subject_order_map.get(s, 999999), s.name)
    )

    final_questions_list = []
    for subject in sorted_subjects:
        subject_questions = questions_by_subject[subject]
        # 5. Сортируем вопросы ВНУТРИ предмета
        subject_questions.sort(key=lambda q: order_map.get(q.id, 999999))
        final_questions_list.extend(subject_questions)

    # Добавляем варианты ответов
    for q in final_questions_list:
        q.fixed_options = q.options.all().order_by('order', 'id')

    context = {
        'test': test,
        'all_questions': final_questions_list, 
        'header_left': f'СИНФИ {test.school_class.name}',
        'header_center': f'ТЕСТИ УМУМӢ {test.test_number}',
        'header_right': test.quarter.year.name if test.quarter and test.quarter.year else str(test.test_date.year),
        'page_title': f"Предпросмотр: {test.name}"
    }
    return render(request, 'booklet/booklet.html', context)


@login_required
@require_POST
def save_booklet_order(request, test_pk):
    if not request.user.is_staff: return JsonResponse({'status': 'error', 'message': 'Нет доступа'}, status=403)
    try:
        test = get_object_or_404(GatTest, pk=test_pk)
        data = json.loads(request.body)
        question_ids = data.get('order')
        if not isinstance(question_ids, list): return JsonResponse({'status': 'error', 'message': 'Неверный формат данных'}, status=400)
        valid_ids = [int(id_str) for id_str in question_ids if str(id_str).isdigit()]
        test.question_order = valid_ids
        test.save()
        return JsonResponse({'status': 'success', 'message': 'Порядок сохранен'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


# =============================================================================
# --- EXPORT TO WORD (DOCX) ---
# =============================================================================

@login_required
def export_booklet_docx(request, test_pk):
    """
    Экспорт теста в MS Word с использованием секционных разрывов.
    Создает структуру: Шапка (1 колонка) -> Разрыв -> Вопросы (2 колонки).
    """
    test = get_object_or_404(GatTest, pk=test_pk)
    
    doc = Document()
    
    # --- НАСТРОЙКА СТИЛЕЙ ---
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    # --- СЕКЦИЯ 1: ШАПКА (1 колонка, на всю ширину) ---
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(12.7) # Узкие поля (0.5 дюйма)
    section.right_margin = Mm(12.7)
    section.top_margin = Mm(12.7)
    section.bottom_margin = Mm(12.7)
    
    # Текст шапки
    header_text = f"СИНФИ {test.school_class.name}   |   ТЕСТИ УМУМӢ {test.test_number}   |   {test.test_date.year}"
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run(header_text)
    run.bold = True
    run.font.size = Pt(12)
    
    # Линия разделитель (имитация через подчеркивание)
    doc.add_paragraph( "_" * 85 ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph() # Пустая строка отступа

    # --- СЕКЦИЯ 2: ВОПРОСЫ (2 колонки) ---
    # Добавляем "Непрерывный разрыв раздела" (Continuous Section Break)
    new_section = doc.add_section(WD_SECTION.CONTINUOUS)
    
    # Применяем магию XML для включения 2-х колонок в этой секции
    set_columns(new_section, 2)

    # --- ПОДГОТОВКА ДАННЫХ (СОРТИРОВКА) ---
    raw_questions = list(test.questions.all())
    order_map = {qid: idx for idx, qid in enumerate(test.question_order or [])}
    
    from collections import defaultdict
    questions_by_subject = defaultdict(list)
    for q in raw_questions:
        questions_by_subject[q.subject].append(q)

    subject_order_map = {}
    for subject, q_list in questions_by_subject.items():
        if q_list:
            min_index = min([order_map.get(q.id, 999999) for q in q_list])
        else:
            min_index = 999999
        subject_order_map[subject] = min_index

    sorted_subjects = sorted(
        questions_by_subject.keys(),
        key=lambda s: (subject_order_map.get(s, 999999), s.name)
    )

    q_counter = 1

    # --- ГЕНЕРАЦИЯ КОНТЕНТА ВОПРОСОВ ---
    for subject in sorted_subjects:
        # 1. Заголовок предмета
        p_subj = doc.add_paragraph()
        p_subj.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_subj = p_subj.add_run(subject.name.upper())
        run_subj.bold = True
        run_subj.font.size = Pt(12)
        p_subj.paragraph_format.space_before = Pt(12)
        p_subj.paragraph_format.space_after = Pt(6)
        p_subj.paragraph_format.keep_with_next = True # Приклеить к следующему

        subject_questions = questions_by_subject[subject]
        subject_questions.sort(key=lambda q: order_map.get(q.id, 999999))

        for q in subject_questions:
            # 2. Текст вопроса
            p_q = doc.add_paragraph()
            p_q.paragraph_format.keep_together = True     # Не разрывать абзац
            p_q.paragraph_format.keep_with_next = True    # Не отрывать от картинки/ответов
            
            run_num = p_q.add_run(f"{q_counter}. ")
            run_num.bold = True
            p_q.add_run(q.text)
            
            # 3. Картинка (если есть)
            if q.question_image:
                try:
                    doc.add_picture(q.question_image.path, width=Mm(80)) # Ширина под колонку
                    last_p = doc.paragraphs[-1]
                    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    last_p.paragraph_format.keep_with_next = True
                except Exception:
                    pass

            # 4. Варианты ответов
            options = q.options.all().order_by('order', 'id')
            
            if options.count() == 4:
                # Таблица 2x2 для аккуратности
                table = doc.add_table(rows=2, cols=2)
                table.autofit = False 
                # Заполняем
                letters = ['A', 'B', 'C', 'D']
                for i, opt in enumerate(options):
                    row = i // 2
                    col = i % 2
                    cell = table.cell(row, col)
                    p_opt = cell.paragraphs[0]
                    p_opt.add_run(f"{letters[i]}) ").bold = True
                    p_opt.add_run(opt.text)
                    
                # Небольшой отступ после таблицы
                doc.add_paragraph().paragraph_format.space_after = Pt(2)

            else:
                # Список, если не 4 варианта
                for i, opt in enumerate(options):
                    letter = chr(65 + i)
                    p_opt = doc.add_paragraph(f"{letter}) {opt.text}")
                    p_opt.paragraph_format.left_indent = Mm(5)
                    p_opt.paragraph_format.space_after = Pt(0)

            q_counter += 1
            
            # Пустая строка между вопросами
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(6)

    # --- ОТПРАВКА ФАЙЛА ---
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=booklet_{test.pk}.docx'
    doc.save(response)
    return response

@login_required
def download_booklet_pdf(request, test_pk):
    """
    Генерирует PDF на сервере и отдает его пользователю.
    """
    # 1. Получаем данные (точно так же, как в preview)
    test = get_object_or_404(
        GatTest.objects.prefetch_related('questions__options', 'questions__subject'), 
        pk=test_pk
    )
    
    raw_questions = list(test.questions.all())
    order_map = {qid: idx for idx, qid in enumerate(test.question_order or [])}

    from collections import defaultdict
    questions_by_subject = defaultdict(list)
    for q in raw_questions:
        questions_by_subject[q.subject].append(q)

    subject_order_map = {}
    for subject, q_list in questions_by_subject.items():
        if q_list:
            min_index = min([order_map.get(q.id, 999999) for q in q_list])
        else:
            min_index = 999999
        subject_order_map[subject] = min_index

    sorted_subjects = sorted(
        questions_by_subject.keys(),
        key=lambda s: (subject_order_map.get(s, 999999), s.name)
    )

    final_questions_list = []
    for subject in sorted_subjects:
        subject_questions = questions_by_subject[subject]
        subject_questions.sort(key=lambda q: order_map.get(q.id, 999999))
        final_questions_list.extend(subject_questions)

    for q in final_questions_list:
        q.fixed_options = q.options.all().order_by('order', 'id')

    # 2. Готовим контекст
    context = {
        'test': test,
        'all_questions': final_questions_list, 
        'header_left': f'СИНФИ {test.school_class.name}',
        'header_center': f'ТЕСТИ УМУМӢ {test.test_number}',
        'header_right': str(test.test_date.year),
        # Флаг для шаблона, что мы в режиме PDF (чтобы скрыть лишнее)
        'is_pdf_mode': True 
    }

    # 3. Рендерим HTML в строку
    # Важно: создайте отдельный чистый шаблон для PDF или используйте существующий с условиями
    html_string = render_to_string('booklet/booklet_pdf.html', context, request=request)

    # 4. Превращаем HTML в PDF с помощью WeasyPrint
    # base_url нужен, чтобы WeasyPrint нашел картинки и стили на диске
    pdf_file = weasyprint.HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf()

    # 5. Отдаем файл
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="booklet_{test.pk}.pdf"'
    return response